#!/usr/bin/env python3
"""Create a reference document for pandoc with proper styling."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Get the default styles
styles = doc.styles

# Configure Normal (Body Text) style
normal_style = styles['Normal']
normal_font = normal_style.font
normal_font.name = 'Verdana'
normal_font.size = Pt(9)
normal_para = normal_style.paragraph_format
normal_para.space_before = Pt(0)
normal_para.space_after = Pt(12)
normal_para.line_spacing = 1.4

# Configure Heading 1 style
h1_style = styles['Heading 1']
h1_font = h1_style.font
h1_font.name = 'Verdana'
h1_font.size = Pt(9)
h1_font.bold = True
h1_font.color.rgb = RGBColor(0, 0, 0)
h1_para = h1_style.paragraph_format
h1_para.space_before = Pt(0)
h1_para.space_after = Pt(12)
h1_para.line_spacing = 1.4

# Configure Heading 2 style (with numbering)
h2_style = styles['Heading 2']
h2_font = h2_style.font
h2_font.name = 'Verdana'
h2_font.size = Pt(9)
h2_font.bold = True
h2_font.color.rgb = RGBColor(0, 0, 0)
h2_para = h2_style.paragraph_format
h2_para.space_before = Pt(12)
h2_para.space_after = Pt(12)
h2_para.line_spacing = 1.4

# Configure Heading 3 style
h3_style = styles['Heading 3']
h3_font = h3_style.font
h3_font.name = 'Verdana'
h3_font.size = Pt(9)
h3_font.bold = True
h3_font.color.rgb = RGBColor(0, 0, 0)
h3_para = h3_style.paragraph_format
h3_para.space_before = Pt(12)
h3_para.space_after = Pt(12)
h3_para.line_spacing = 1.4

# Add sample content to demonstrate the styles
doc.add_heading('Sample Heading 1', level=1)
doc.add_paragraph('This is a sample paragraph with Verdana 9pt font.')
doc.add_heading('Sample Heading 2', level=2)
doc.add_paragraph('This is another sample paragraph.')

# Save the reference document
doc.save('templates/reference.docx')
print('Reference document created: templates/reference.docx')
